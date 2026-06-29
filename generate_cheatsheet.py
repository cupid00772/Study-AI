from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
# set extremely narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Add title
title = doc.add_heading('機率與統計 考前終極小抄 (含所有連續分配與CLT)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.style.font.size = Pt(14)

def add_heading(text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    run = h.runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text)

def add_weakness(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("⚠️ 陷阱: ")
    r.bold = True
    r.font.color.rgb = RGBColor(204, 0, 0)
    p.add_run(text)

add_heading("Ch 2 機率與貝氏定理")
add_bullet("條件機率: P(A | B) = P(A ∩ B) / P(B) | 獨立: P(A ∩ B) = P(A)P(B) (互斥為 P(A ∩ B)=0，兩者不同)")
add_bullet("全機率: P(B) = Σ P(B|A_i)P(A_i) | 貝氏定理: P(A_i|B) = [P(B|A_i)P(A_i)] / P(B)")
add_weakness("抽出不放回注意分母遞減。瑕疵分類題必畫樹狀圖。")

add_heading("Ch 3 離散型分配 (Discrete)")
add_bullet("Bernoulli(p): P(1)=p。E=p, Var=p(1-p) | Binomial(n,p): n次中k次成功。E=np, Var=np(1-p)")
add_bullet("Poisson(λ): P(k)=(e^-λ * λ^k)/k!。E=λ, Var=λ (時間若改變，λ要按比例縮放)")
add_bullet("Hypergeometric(N,K,n): 抽出不放回。P(k)=[C(K,k)*C(N-K,n-k)] / C(N,n)")
add_bullet("Negative Binomial(r,p): 直到第r次成功共需n次。P(n)=C(n-1, r-1) * p^r * (1-p)^(n-r)。E=r/p")

add_heading("Ch 4 & 9 連續型分配 (Continuous)")
add_bullet("均勻 Uniform(a,b): f(x)=1/(b-a), E=(a+b)/2, Var=(b-a)²/12")
add_bullet("常態 Normal(μ,σ²): 鐘型對稱。68-95-99.7% 法則 (μ±1σ, μ±2σ, μ±3σ)。標準化 Z = (X-μ)/σ")
add_bullet("對數常態 Lognormal: 若 X~N(μ,σ²), 則 Y=e^X 為對數常態。")
add_bullet("指數 Exponential(λ): 等待時間。f(x)=λe^(-λx), E=1/λ, Var=1/λ²。(具備 Lack of Memory 失憶性: P(T>t+s|T>s)=P(T>t))")
add_bullet("Gamma(r,λ): 等待第r次發生的時間。E=r/λ, Var=r/λ² | Weibull: 常估算元件壽命。")
add_weakness("求條件期望值 E[X | X≤a] 必除以 P(X≤a)。常態尾部機率(如大於μ+2σ)是 (100%-95%)/2 = 2.5%。")

add_heading("Ch 6 聯合機率 & Ch 8 測量誤差與點估計")
add_bullet("邊際機率: 從表格中把「同一列」或「同一欄」相加。獨立性: 每個格子皆滿足 P(x,y)=P(x)*P(y)")
add_bullet("期望值/變異數線性組合: E[aX±bY+c] = aE[X]±bE[Y]+c | 若獨立，Var(aX±bY+c) = a²Var(X) + b²Var(Y)")
add_bullet("誤差模型: Measured = True + Bias(系統誤差, 固定) + Random Error(隨機誤差, E=0)")
add_bullet("點估計 MSE (Mean Squared Error) = Var(θ) + [Bias(θ)]²。樣本平均是Unbiased，樣本變異數是Asymptotically Unbiased。")

add_heading("中央極限定理 (CLT) & 常態近似")
add_bullet("CLT: 若 n>30，樣本平均 X̄ 近似常態 N(μ, σ²/n)；樣本總和 S_n 近似 N(nμ, nσ²)。")
add_bullet("近似 Binomial: 若 np>10 且 n(1-p)>10，則 Binomial(n,p) 近似 N(np, np(1-p))。")
add_bullet("近似 Poisson: 若 λ>10，則 Poisson(λ) 近似 N(λ, λ)。")
add_weakness("連續校正 (Continuity Correction): 離散轉連續時，端點要 ±0.5 (例如求 P(X≤45) 要改算 P(X<45.5))。")

doc.save('機率與統計小抄.docx')
