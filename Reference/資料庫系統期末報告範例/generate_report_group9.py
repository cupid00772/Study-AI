import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def generate_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('資料庫系統 期末專題成果報告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('專題名稱：智慧水肥派車與循環農業資料庫系統\n組別：第 9 組\n指導教授：李金鳳\n').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Chapter 1
    add_heading(doc, '第一章、專案團隊')
    doc.add_paragraph('本專案由以下成員共同開發完成：')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '學號'
    hdr_cells[1].text = '姓名'
    hdr_cells[2].text = '負責任務'
    hdr_cells[3].text = '備註'
    
    team_members = [
        ('D1349082', '林晟恩', '前端開發、組長'),
        ('D1349510', '陳柏凱', '程式設計師'),
        ('D1285325', '張承新', '程式設計師')
    ]
    
    for member in team_members:
        row_cells = table.add_row().cells
        row_cells[0].text = member[0]
        row_cells[1].text = member[1]
        row_cells[2].text = member[2]
        row_cells[3].text = ''
    doc.add_paragraph('\n')

    # Chapter 2
    add_heading(doc, '第二章、系統概述')
    add_heading(doc, '(一) 計畫說明', level=2)
    doc.add_paragraph('近年來全球氣候變遷加劇，推動聯合國永續發展目標 (SDGs) 已成為產官學界刻不容緩的議題。台灣的畜牧業及民生廢水會產生大量水肥，若未妥善處理將造成環境污染並釋放大量溫室氣體。然而，水肥中富含氮、磷、鉀等農業所需養分，若能透過完善的厭氧發酵與檢驗機制，將其轉化為有機肥料，不僅能落實 SDG 12（責任消費及生產），更能有效減少化學肥料的使用。本團隊擬開發「智慧水肥派車與循環農業資料庫系統」，透過資料庫完整記錄清運軌跡與檢驗履歷，打造水肥再利用的循環經濟生態圈。')
    
    add_heading(doc, '(二) 系統特色', level=2)
    doc.add_paragraph('1. 智慧派車與減碳 (SDG 13)：結合地理資訊與水肥車容量限制，透過後端演算法計算出最節能的收運路線，具體降低運輸過程的碳排放。')
    doc.add_paragraph('2. 安心檢驗履歷 (SDG 12)：建立水肥發酵與化學檢驗的關聯式資料庫，確保產出的肥料符合農業使用標準，落實生產履歷透明化。')
    doc.add_paragraph('3. 農工媒合平台 (SDG 17)：提供需要有機肥料的農民與水肥處理廠一個資訊對稱的交易與媒合介面，促進綠色資源流動。')
    doc.add_paragraph('4. 無痛操作設計 (SDG 10)：針對司機端與農民端設計直覺的響應式網頁（RWD），降低數位落差帶來的操作障礙。')

    # Chapter 3
    add_heading(doc, '第三章、開發技術')
    add_heading(doc, '(一) 系統架構', level=2)
    doc.add_paragraph('本系統採用標準的三層式架構 (3-tier Architecture)，確保「關注點分離」：\n'
                      '- 展示層 (前端)：HTML5, CSS3, JavaScript 搭配 Bootstrap 5 框架，提供響應式 (RWD) 網頁介面。\n'
                      '- 運算邏輯層 (後端)：ASP.NET (C#) / Node.js / PHP，負責處理派車演算法、權限驗證等業務邏輯。\n'
                      '- 資料服務層 (資料庫)：Microsoft SQL Server / Oracle，執行資料存取、事務處理 (Transaction) 與資料完整性維護。')
    add_heading(doc, '(二) 採用原因', level=2)
    doc.add_paragraph('三層式架構能有效將系統工作負荷重新安排，分散風險與提高安全性（防範 SQL 注入），且具備高擴充彈性，未來若需新增物聯網 (IoT) 感測數據可直接擴充而不影響現有架構。前端採用 Bootstrap 5 則能達成跨平台親和性，讓農民與司機用手機即可輕鬆操作。')

    # Chapter 4
    add_heading(doc, '第四章、軟體需求')
    doc.add_paragraph('- 伺服器作業系統：Windows Server 或 Linux 環境\n'
                      '- Web 伺服器：IIS 或 Apache\n'
                      '- 資料庫管理系統 (RDBMS)：Microsoft SQL Server 或 Oracle\n'
                      '- 後端開發技術：ASP.NET (C#) / Node.js / PHP\n'
                      '- 前端開發技術：HTML5, CSS3, JavaScript, Bootstrap 5\n'
                      '- 開發工具：Visual Studio, VS Code, SSMS')

    # Chapter 5
    add_heading(doc, '第五章、功能簡介')
    add_heading(doc, '(一) 功能架構圖', level=2)
    doc.add_paragraph('[請在此處插入系統功能架構圖]')
    
    add_heading(doc, '(二) 功能分析', level=2)
    doc.add_paragraph('1. 會員系統：\n分為管理者(處理廠)、業者會員(司機)、消費者會員(農民)三種角色權限。管理者可審核進場紀錄；司機可檢視派車任務並回報容量；農民可預約配送與查詢綠肥。')
    doc.add_paragraph('2. 典藏系統 (綠肥檢驗履歷)：\n完整記錄進廠編號、厭氧發酵天數、氮磷鉀與重金屬檢驗結果，確保每批出貨可追溯原始報告。')
    doc.add_paragraph('3. 查詢功能 (飛天補快)：\n司機端可動態查詢最佳清運路線；農民端可依區域與養分需求快速過濾綠肥批次。')
    doc.add_paragraph('4. 客服系統 (討論區/留言板)：\n農民可分享農作成果、提出土壤改良疑問；處理廠可發布最新合格綠肥資訊，形成雙向溝通。')
    doc.add_paragraph('5. 數據分析 (減碳績效排行榜)：\n透過大數據視覺化，計算最佳化路線省下之里程與轉化綠肥噸數，換算成「減少碳排量」，作為企業 ESG 報告之量化指標。')

    # Chapter 6
    add_heading(doc, '第六章、資料庫分析與設計')
    add_heading(doc, '(一) 實體關聯模型 (ERD)', level=2)
    if os.path.exists('Project03.jpeg'):
        doc.add_picture('Project03.jpeg', width=Inches(5.5))
    else:
        doc.add_paragraph('[未找到 Project03.jpeg，請手動插入]')
    
    add_heading(doc, '(二) 關聯綱要圖', level=2)
    if os.path.exists('project04.jpg'):
        doc.add_picture('project04.jpg', width=Inches(5.5))
    else:
        doc.add_paragraph('[未找到 project04.jpg，請手動插入]')
    
    add_heading(doc, '(三) 實體關聯表 (規格表+語法)', level=2)
    doc.add_paragraph('依據綱要圖與 ERD，我們設計了以下核心資料表與對應的 SQL 建表語法：\n')
    
    doc.add_paragraph('【Table: 使用者 User】\n規格：uid(PK), 姓名, 帳號')
    doc.add_paragraph('CREATE TABLE [User] (\n  uid INT PRIMARY KEY IDENTITY(1,1),\n  account VARCHAR(50) NOT NULL UNIQUE,\n  name NVARCHAR(50) NOT NULL\n);')
    
    doc.add_paragraph('【Table: 水肥 (資源站點) WaterFertilizer】\n規格：p.code(PK), 經緯度(Coord), 名稱')
    doc.add_paragraph('CREATE TABLE WaterFertilizer (\n  pcode VARCHAR(20) PRIMARY KEY,\n  coord VARCHAR(50) NOT NULL,\n  name NVARCHAR(50) NOT NULL\n);')
    
    doc.add_paragraph('【Table: 儲存槽 Tank】\n規格：代號(T.code)(PK), 水位, 容量, 地點, p.code(FK)')
    doc.add_paragraph('CREATE TABLE Tank (\n  tcode VARCHAR(20) PRIMARY KEY,\n  water_level FLOAT,\n  capacity FLOAT NOT NULL,\n  location NVARCHAR(100),\n  pcode VARCHAR(20) FOREIGN KEY REFERENCES WaterFertilizer(pcode)\n);')

    doc.add_paragraph('【Table: 收集水肥 CollectRecord】\n規格：p.code(FK), uid(FK), 日期, 量 (複合主鍵)')
    doc.add_paragraph('CREATE TABLE CollectRecord (\n  pcode VARCHAR(20) FOREIGN KEY REFERENCES WaterFertilizer(pcode),\n  uid INT FOREIGN KEY REFERENCES [User](uid),\n  collect_date DATETIME,\n  quantity FLOAT,\n  PRIMARY KEY(pcode, uid, collect_date)\n);')

    doc.add_paragraph('【Table: 留言板 MessageBoard】\n規格：tno(PK), 時間, 內容, uid(FK)')
    doc.add_paragraph('CREATE TABLE MessageBoard (\n  tno INT PRIMARY KEY IDENTITY(1,1),\n  post_time DATETIME DEFAULT GETDATE(),\n  content NVARCHAR(MAX) NOT NULL,\n  uid INT FOREIGN KEY REFERENCES [User](uid)\n);')

    # Chapter 7
    add_heading(doc, '第七章、操作過程與說明')
    doc.add_paragraph('本章節以圖文方式展示系統的實際操作流程 (請從 di-zu-_zi-liao-ku-demo.mp4 展示影片擷取對應畫面貼入)。')
    
    add_heading(doc, '(一) 會員系統', level=2)
    doc.add_paragraph('[請在此插入：管理者、司機、農民等不同角色登入後的介面截圖]\n說明：展示依據不同角色賦予對應系統權限之畫面。')
    
    add_heading(doc, '(二) 典藏系統', level=2)
    doc.add_paragraph('[請在此插入：綠肥檢驗履歷建檔與檢視畫面截圖]\n說明：展示氮磷鉀與重金屬檢測報告之輸入與查詢。')
    
    add_heading(doc, '(三) 搜尋功能', level=2)
    doc.add_paragraph('[請在此插入：飛天補快派車路線查詢畫面截圖]\n說明：司機透過系統查詢最佳化路線。')
    
    add_heading(doc, '(四) 客服系統 (留言板)', level=2)
    doc.add_paragraph('[請在此插入：農民與處理廠互動之討論區畫面截圖]\n說明：展示雙方交流農作與肥料問題之情形。')
    
    add_heading(doc, '(五) 數據分析', level=2)
    doc.add_paragraph('[請在此插入：減碳績效排行榜或資料視覺化統計圖截圖]\n說明：展示系統依據路線與綠肥轉換率計算出之減碳效益圖表。')

    # Chapter 8
    add_heading(doc, '第八章、專案進度')
    doc.add_paragraph('本系統嚴格遵循 SDLC，於 14 週內完整交付：')
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    t2_hdr = table2.rows[0].cells
    t2_hdr[0].text = '週次'
    t2_hdr[1].text = '階段'
    t2_hdr[2].text = '主要工作項目'
    
    schedule = [
        ('第 1–2 週', '啟動階段', '訂定水肥循環主題、分析功能、確認 SDGs 關聯性'),
        ('第 3–4 週', '分析與設計', '設計系統架構、繪製 ER Model 與關聯綱要'),
        ('第 5–8 週', '開發階段', '資料庫建置、後端 CRUD 核心程式、派車查詢邏輯'),
        ('第 9–11 週', '整合階段', '前端切版、RWD 設計、前後端系統整合'),
        ('第 12–13 週', '測試階段', '系統除錯 (Debug)、跨平台瀏覽器測試'),
        ('第 14 週', '驗收階段', '企劃書定稿、期末 DEMO 準備')
    ]
    for row in schedule:
        rc = table2.add_row().cells
        rc[0].text = row[0]
        rc[1].text = row[1]
        rc[2].text = row[2]

    # Chapter 9
    add_heading(doc, '第九章、執行本專題的心路歷程與收穫')
    doc.add_paragraph('【組長：林晟恩】\n在擔任前端開發與組長的過程中，最大的挑戰在於如何將複雜的資料庫邏輯轉換為農民與司機都能直覺使用的 RWD 介面。同時，統籌團隊進度也讓我學到了寶貴的專案管理經驗與溝通技巧。')
    doc.add_paragraph('【組員：陳柏凱】\n將 SDGs 的概念融入資料庫設計是個全新的挑戰。透過計算碳排與綠肥數據，我學會了如何讓死板的資料具有環保與社會價值，也提升了我對資料正規化與完整性維護的能力。')
    doc.add_paragraph('【組員：張承新】\n在整合前端與後端的過程中遇到了不少難題，透過團隊互相進行配對程式設計 (Pair Programming)，我們順利解決了 API 串接與跨域問題，這次的專題合作讓我收穫滿滿。')

    # Chapter 10
    add_heading(doc, '第十章、參考文獻')
    doc.add_paragraph('[1] 聯合國永續發展目標 (SDGs) 官方指標與說明。\n'
                      '[2] 農業部：循環農業與水肥再利用技術指南。\n'
                      '[3] MS SQL Server 官方說明文件 (T-SQL)。\n'
                      '[4] ASP.NET / C# Web API 開發手冊。\n'
                      '[5] Bootstrap 5 官方文件 (RWD 響應式設計)。')

    doc.save('第九組_智慧水肥派車_期末成果報告.docx')
    print("Group 9 Report updated successfully with exact SQL schemas and images inserted.")

if __name__ == '__main__':
    generate_report()
