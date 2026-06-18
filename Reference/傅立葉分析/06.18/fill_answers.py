import docx
from docx.shared import Inches

doc_path = r"c:\Users\cupid\Downloads\Stuty-space\Reference\傅立葉分析\06.18\課堂學習單_06.18.docx"
out_path = r"c:\Users\cupid\Downloads\Stuty-space\Reference\傅立葉分析\06.18\課堂學習單_06.18_回答.docx"

doc = docx.Document(doc_path)

answers = {
    "A.": [
        "二維 (2D)。",
        "空間頻率 (Spatial frequency)。",
        "低頻 (經過 fftshift 後，頻譜圖中心點代表直流與低頻成分)。",
        "影像的邊緣、輪廓、細節特徵或雜訊。",
        "濾除高頻成分，保留低頻成分，達到影像平滑化與降低雜訊的目的。"
    ],
    "B.": [
        "D0 越小的影像最模糊 (由 Matlab 結果可見 D0=10 最模糊)。因為 D0 越小代表濾除的高頻成分(即邊緣與細節)越多，只剩下低頻成分，導致影像整體變得平滑且模糊。",
        "D0 越大的影像 (如 D0=150) 最接近原始影像，因為保留了大部分的高頻與低頻成分。",
        "應選擇中等大小的 D0 (例如 D0=80)，這樣可濾除部分極高頻的雜訊，同時保留足夠的邊緣資訊以維持人物輪廓清晰。",
        "high-pass filter (高通濾波器)。",
        "否。過度平滑(過於模糊)會導致影像中重要的邊緣、輪廓和細節特徵丟失，使得影像內容難以辨識。",
        "不適合。因為文字的邊界輪廓屬於高頻成分，若使用很小的 D0，會將這些高頻成分濾除，導致文字模糊到無法閱讀。",
        "應選擇一個適當大小 (折衷) 的 D0。該 D0 需夠小以濾掉最高頻的雜訊，但又需夠大以保留構成邊緣的中高頻成分。通常需要透過實驗(如 Matlab 測試)找出最佳的 D0。"
    ]
}

a_idx = 0
b_idx = 0
current_section = None

for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith("A."):
        current_section = "A."
    elif text.startswith("B."):
        current_section = "B."
        
    if text.startswith("答："):
        if current_section == "A." and a_idx < len(answers["A."]):
            para.text = "答：" + answers["A."][a_idx]
            a_idx += 1
        elif current_section == "B." and b_idx < len(answers["B."]):
            para.text = "答：" + answers["B."][b_idx]
            b_idx += 1

doc.add_page_break()
doc.add_heading("Matlab FFT 處理結果", level=1)
doc.add_paragraph("以下為使用 Matlab 對提供的圖片進行 2D FFT 以及不同 D0 截止頻率下的 Ideal Low Pass Filter (ILPF) 處理結果：")

doc.add_heading("圖片1: S__11788318乙.jpg", level=2)
doc.add_picture(r"c:\Users\cupid\Downloads\Stuty-space\Reference\傅立葉分析\06.18\img1_analysis.png", width=Inches(6.0))

doc.add_heading("圖片2: S__11878405乙.jpg", level=2)
doc.add_picture(r"c:\Users\cupid\Downloads\Stuty-space\Reference\傅立葉分析\06.18\img2_analysis.png", width=Inches(6.0))

doc.save(out_path)
print("Saved to", out_path)
