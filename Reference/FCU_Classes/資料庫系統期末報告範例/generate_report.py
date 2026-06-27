import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def generate_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('資料庫系統 期末專題成果報告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('專題名稱：綠能查詢網站 (ECO friendly)\n組別：第八組 (暫定)\n').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Chapter 1
    add_heading(doc, '第一章、專案團隊')
    doc.add_paragraph('本專案由以下成員共同開發完成：')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '學號'
    hdr_cells[1].text = '姓名'
    hdr_cells[2].text = '負責任務'
    hdr_cells[3].text = '備註'
    
    # Placeholder rows
    for _ in range(3):
        row_cells = table.add_row().cells
        row_cells[0].text = '[學號]'
        row_cells[1].text = '[姓名]'
        row_cells[2].text = '[如：前端介面設計、資料庫建置]'
        row_cells[3].text = ''
    doc.add_paragraph('\n')

    # Chapter 2
    add_heading(doc, '第二章、系統概述')
    add_heading(doc, '(一) 計畫說明', level=2)
    doc.add_paragraph('隨著全球氣候變遷加劇與環保意識抬頭，再生能源(綠能)的發展與應用成為各國關注的焦點。本專題旨在建立一個「綠能查詢網站 (ECO friendly)」，提供使用者查詢各種綠能發電廠（如風力發電、太陽能發電、火力發電等）的相關資訊，並具備會員機制、資料典藏與留言板互動功能，讓社會大眾能更輕鬆地了解台灣各地的發電設施與綠能現況。')
    
    add_heading(doc, '(二) 系統特色', level=2)
    doc.add_paragraph('1. 豐富的發電廠資訊：涵蓋全台多處太陽能、風力等發電站資料。\n'
                      '2. 會員專屬功能：提供會員註冊與登入，登入後可進行資料收藏與評論。\n'
                      '3. 響應式網頁設計：前端採用 HTML/CSS/JavaScript，提供流暢的圖片輪播與使用者介面。\n'
                      '4. 互動留言板：使用者可對特定的綠能設施發表看法，促進資訊交流。')

    # Chapter 3
    add_heading(doc, '第三章、開發技術')
    add_heading(doc, '(一) 系統架構', level=2)
    doc.add_paragraph('本系統採用三層式架構 (3-tier Architecture)：\n'
                      '- 展示層 (前端)：HTML, CSS, JavaScript (jQuery), 提供響應式且直覺的操作介面。\n'
                      '- 應用層 (後端)：PHP，負責處理使用者請求、業務邏輯及資料處理。\n'
                      '- 資料層 (資料庫)：MySQL (MariaDB)，負責儲存會員、發電廠、留言等關聯式資料。')
    add_heading(doc, '(二) 採用原因', level=2)
    doc.add_paragraph('選擇 PHP 搭配 MySQL 的 LAMP/WAMP 架構，是因為其開源且社群資源豐富，適合快速開發 Web 應用程式。前端引入 jQuery 與自定義 CSS 可大幅降低開發成本並提升網頁視覺效果。')

    # Chapter 4
    add_heading(doc, '第四章、軟體需求')
    doc.add_paragraph('- 作業系統：Windows 10 / 11 或 macOS\n'
                      '- 伺服器環境：XAMPP (包含 Apache 與 MariaDB)\n'
                      '- 資料庫系統：MySQL / MariaDB\n'
                      '- 後端語言：PHP 8.x\n'
                      '- 前端技術：HTML5, CSS3, JavaScript (jQuery 3.6.0)\n'
                      '- 開發工具：Visual Studio Code')

    # Chapter 5
    add_heading(doc, '第五章、功能簡介')
    add_heading(doc, '(一) 功能架構圖', level=2)
    doc.add_paragraph('[請在此處插入系統功能架構圖]')
    
    add_heading(doc, '(二) 功能分析', level=2)
    doc.add_paragraph('1. 會員系統：\n提供使用者註冊 (regist.php) 與登入 (login.php) 功能，密碼具備安全驗證，保護會員隱私。')
    doc.add_paragraph('2. 典藏系統：\n會員可將感興趣的發電廠或綠能資訊加入「我的收藏」(collections.php)，並支援新增與刪除收藏 (remove_collection.php)。')
    doc.add_paragraph('3. 查詢功能：\n提供資料庫查詢介面 (search.php)，可依據發電類型（如太陽能、風力）或地區查詢相關發電廠資訊。')
    doc.add_paragraph('4. 客服系統 (留言板)：\n提供公開的留言討論區 (board.php)，讓使用者分享對綠電發展的看法。')
    doc.add_paragraph('5. 數據分析：\n系統後台收集了各發電站的發電量 (Dynamo) 數據，可作為數據分析之基礎，呈現各類綠能發電的貢獻度排行。')

    # Chapter 6
    add_heading(doc, '第六章、資料庫分析與設計')
    add_heading(doc, '(一) 實體關聯模型 (ERD)', level=2)
    doc.add_paragraph('[請在此插入 ERD 圖片 (例如 Project03.jpeg)]')
    
    add_heading(doc, '(二) 關聯綱要圖', level=2)
    doc.add_paragraph('[請在此插入綱要圖 (例如 project04.jpg)]')
    
    add_heading(doc, '(三) 實體關聯表 (規格表+語法)', level=2)
    doc.add_paragraph('本系統包含多個資料表，以下列舉重要資料表規格：\n')
    
    doc.add_paragraph('【Table: users (會員資料表)】')
    doc.add_paragraph('CREATE TABLE `users` (\n  `u_no` int(11) NOT NULL AUTO_INCREMENT,\n  `u_account` varchar(30) NOT NULL,\n  `u_password` varchar(100) NOT NULL,\n  `u_name` varchar(30) NOT NULL,\n  `u_address` varchar(200) DEFAULT NULL,\n  PRIMARY KEY (`u_no`),\n  UNIQUE KEY `u_account` (`u_account`)\n);')
    
    doc.add_paragraph('【Table: dyna (發電廠資料表)】')
    doc.add_paragraph('CREATE TABLE `dyna` (\n  `d_id` int(11) NOT NULL AUTO_INCREMENT,\n  `address` varchar(50) DEFAULT NULL,\n  `type` varchar(50) DEFAULT NULL,\n  `Dynamo` int(11) DEFAULT NULL,\n  `d_name` varchar(15) NOT NULL,\n  PRIMARY KEY (`d_id`)\n);')
    
    doc.add_paragraph('【Table: messages (留言板資料表)】')
    doc.add_paragraph('CREATE TABLE `messages` (\n  `m_no` int(11) NOT NULL AUTO_INCREMENT,\n  `m_name` varchar(50) DEFAULT NULL,\n  `m_content` varchar(200) DEFAULT NULL,\n  `m_date` datetime DEFAULT NULL,\n  `u_no` int(11) DEFAULT NULL,\n  PRIMARY KEY (`m_no`),\n  FOREIGN KEY (`u_no`) REFERENCES `users` (`u_no`)\n);')
    
    # Chapter 7
    add_heading(doc, '第七章、操作過程與說明')
    doc.add_paragraph('本章節以圖文方式展示系統的實際操作流程 (請從展示影片擷取對應畫面貼入)。')
    
    add_heading(doc, '(一) 會員系統', level=2)
    doc.add_paragraph('[請在此插入：會員註冊、登入畫面截圖]\n說明：使用者填寫帳號密碼進行註冊，登入後即可使用進階功能。')
    
    add_heading(doc, '(二) 典藏系統', level=2)
    doc.add_paragraph('[請在此插入：加入收藏、檢視收藏列表畫面截圖]\n說明：在瀏覽發電廠時點擊收藏，並可至個人頁面管理收藏名單。')
    
    add_heading(doc, '(三) 搜尋功能', level=2)
    doc.add_paragraph('[請在此插入：關鍵字搜尋與結果呈現畫面截圖]\n說明：支援輸入發電站名稱或類型進行快速篩選。')
    
    add_heading(doc, '(四) 客服系統 (留言板)', level=2)
    doc.add_paragraph('[請在此插入：新增留言與留言列表畫面截圖]\n說明：使用者輸入名稱與內容發佈留言，系統自動記錄時間。')
    
    add_heading(doc, '(五) 數據分析', level=2)
    doc.add_paragraph('[請在此插入：發電數據或排行榜統計畫面截圖]\n說明：展示各地區或各類型綠能的統計數據。')

    # Chapter 8
    add_heading(doc, '第八章、專案進度')
    doc.add_paragraph('以下為本專題之進度甘特圖與里程碑規劃：\n[請在此處插入進度甘特圖圖片或表格]')

    # Chapter 9
    add_heading(doc, '第九章、執行本專題的心路歷程與收穫')
    doc.add_paragraph('【組員A：姓名】\n(請描述在系統開發過程中遇到的困難，例如前後端整合或資料庫設計，以及最後學到的經驗)')
    doc.add_paragraph('【組員B：姓名】\n(請描述在除錯或 UI 設計時的收穫，以及團隊合作的感想)')

    # Chapter 10
    add_heading(doc, '第十章、參考文獻')
    doc.add_paragraph('[1] PHP 官方手冊: https://www.php.net/manual/zh/\n'
                      '[2] MySQL 官方文件: https://dev.mysql.com/doc/\n'
                      '[3] 台灣電力公司 綠能發展概況: https://www.taipower.com.tw/\n'
                      '[4] W3Schools HTML/CSS/JS 教學: https://www.w3schools.com/')

    doc.save('綠能查詢網站_期末成果報告.docx')
    print("Report generated successfully.")

if __name__ == '__main__':
    generate_report()
